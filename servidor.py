import http.server
import socketserver
import json
import urllib.parse
import os
import re
from pathlib import Path
import pypdf
import docx

PORT = 8000
historico_desfazer = []

import urllib.request
import ssl
import threading

# Caminho do dicionário local e conjunto de palavras carregado
DICT_PATH = os.path.join(os.path.dirname(os.path.abspath(__file__)), "portuguese_words.txt")
WORD_SET = set()

# Lista de palavras comuns para fallback offline imediato
COMMON_PT_WORDS = {
    "contrato", "de", "emprestimo", "trabalho", "servico", "venda", "compra", "aluguel",
    "recibo", "nota", "fiscal", "declaracao", "imposto", "renda", "relatorio", "projeto",
    "fatura", "pagamento", "comprovante", "identidade", "cpf", "rg", "cnh", "certidao",
    "nascimento", "casamento", "obito", "procuracao", "laudo", "medico", "exame",
    "curriculo", "foto", "imagem", "video", "audio", "musica", "aula", "curso",
    "manual", "livro", "artigo", "tcc", "monografia", "dissertacao", "tese",
    "apresentacao", "slide", "tabela", "planilha", "cadastro", "cliente", "fornecedor",
    "funcionario", "empresa", "banco", "extrato", "saldo", "transferencia", "pix",
    "boleto", "seguro", "sinistro", "police", "proposta", "orcamento", "pedido",
    "entrega", "frete", "envio", "recebimento", "devolucao", "troca", "garantia",
    "suporte", "atendimento", "chamado", "reclamacao", "sugestao", "elogio", "pesquisa",
    "satisfacao", "ademicon", "reclame", "aqui", "unidade", "colombo", "hauer",
    "nao", "paga", "comissao", "metas", "abusivas", "rescisao", "desconto", "pj",
    "novo", "terceira", "vez", "falta", "atendimento", "assinatura", "assinado"
}

def carregar_dicionario():
    global WORD_SET
    WORD_SET = set(COMMON_PT_WORDS)
    if os.path.exists(DICT_PATH):
        try:
            with open(DICT_PATH, "r", encoding="utf-8") as f:
                for line in f:
                    word = line.strip().lower()
                    if word:
                        WORD_SET.add(word)
            return
        except Exception:
            pass
            
    # Tenta baixar o dicionário caso não exista localmente
    url = "https://raw.githubusercontent.com/kkrypt0nn/wordlists/main/wordlists/languages/portuguese.txt"
    try:
        ctx = ssl._create_unverified_context()
        with urllib.request.urlopen(url, timeout=5, context=ctx) as response:
            content = response.read().decode("utf-8")
            with open(DICT_PATH, "w", encoding="utf-8") as f:
                f.write(content)
            for line in content.splitlines():
                word = line.strip().lower()
                if word:
                    WORD_SET.add(word)
    except Exception:
        pass

# Inicia o carregamento em background para não bloquear o servidor no startup
threading.Thread(target=carregar_dicionario, daemon=True).start()

def remover_acentos(txt):
    d = {
        'á': 'a', 'à': 'a', 'â': 'a', 'ã': 'a', 'ä': 'a',
        'é': 'e', 'è': 'e', 'ê': 'e', 'ë': 'e',
        'í': 'i', 'ì': 'i', 'î': 'i', 'ï': 'i',
        'ó': 'o', 'ò': 'o', 'ô': 'o', 'õ': 'o', 'ö': 'o',
        'ú': 'u', 'ù': 'u', 'û': 'u', 'ü': 'u',
        'ç': 'c', 'ñ': 'n'
    }
    return "".join(d.get(c, c) for c in txt.lower())

def segmentar_palavras(texto_junto):
    if not texto_junto:
        return ""
        
    texto_limpo = remover_acentos(texto_junto)
    n = len(texto_limpo)
    
    dp = [(float('inf'), [])] * (n + 1)
    dp[0] = (0, [])
    
    short_words = {'a', 'o', 'e', 'de', 'do', 'da', 'em', 'um', 'se', 'no', 'na', 'os', 'as'}
    
    for i in range(1, n + 1):
        for j in range(max(0, i - 15), i):
            palavra = texto_limpo[j:i]
            if palavra in WORD_SET or palavra.isdigit() or (len(palavra) == 1 and palavra in short_words):
                custo = dp[j][0] + 1
                if len(palavra) == 1 and palavra not in short_words:
                    custo += 10
                if custo < dp[i][0]:
                    dp[i] = (custo, dp[j][1] + [texto_junto[j:i]])
                    
    if dp[n][0] != float('inf'):
        return " ".join(dp[n][1])
    else:
        palavras = []
        i = 0
        while i < n:
            match = None
            for j in range(min(n, i + 15), i, -1):
                sub = texto_limpo[i:j]
                if sub in WORD_SET or sub.isdigit() or (len(sub) == 1 and sub in short_words):
                    match = texto_junto[i:j]
                    i = j
                    break
            if match is None:
                palavras.append(texto_junto[i])
                i += 1
            else:
                palavras.append(match)
        return " ".join(palavras)

def tentar_segmentar_palavra(palavra):
    if len(palavra) < 6:
        return palavra
    
    palavra_low = remover_acentos(palavra.lower())
    if palavra_low in WORD_SET:
        return palavra
        
    if not palavra_low.isalpha():
        return palavra
        
    segmentada = segmentar_palavras(palavra)
    partes = segmentada.split()
    if len(partes) <= 1:
        return palavra
        
    partes_validas = 0
    for p in partes:
        p_low = remover_acentos(p.lower())
        if len(p_low) == 1:
            if p_low in {'a', 'o', 'e'}:
                partes_validas += 1
        elif p_low in WORD_SET or p_low in {'de', 'do', 'da', 'em', 'um', 'se', 'no', 'na', 'os', 'as'}:
            partes_validas += 1
            
    if partes_validas / len(partes) >= 0.8:
        return segmentada
    return palavra

def obter_numero_diretorio(nome_dir):
    match = re.match(r'^(\d+)', nome_dir)
    if match:
        return match.group(1)
    return None

def limpar_titulo_para_arquivo(titulo):
    if not titulo:
        return ""
    titulo = re.sub(r'\.pptx(\.pptx)?$', '', titulo, flags=re.IGNORECASE)
    titulo = re.sub(r'\.pdf$', '', titulo, flags=re.IGNORECASE)
    titulo = re.sub(r'^[:\s|]+', '', titulo)
    if " | " in titulo:
        titulo = titulo.split(" | ")[0]
    elif " - " in titulo:
        partes = titulo.split(" - ")
        if len(partes) > 1 and any(x in partes[1].lower() for x in ['plataforma', 'dica', 'unisuam', 'site', 'revista', 'pdf']):
            titulo = partes[0]
    titulo = re.sub(r'[\\/*?:"<>|]', '', titulo)
    
    # Separar palavras em camelCase / PascalCase e substituir vírgulas por espaços
    titulo = titulo.replace(',', ' ')
    titulo = re.sub(r'([a-z0-9])([A-Z])', r'\1 \2', titulo)
    titulo = re.sub(r'([A-Z]+)([A-Z][a-z])', r'\1 \2', titulo)
    
    titulo = titulo.replace('_', ' ').replace('-', ' ').replace('+', ' ')
    titulo = re.sub(r'\s+', ' ', titulo).strip()
    if len(titulo) > 100:
        titulo = titulo[:97] + "..."
        
    palavras = titulo.split()
    preposicoes = {'de', 'do', 'da', 'dos', 'das', 'e', 'a', 'o', 'em', 'para', 'com', 'por', 'ou', 'um', 'uma'}
    resultado = []
    for i, palavra in enumerate(palavras):
        palavra_processada = tentar_segmentar_palavra(palavra)
        for sub_palavra in palavra_processada.split():
            if len(resultado) == 0:
                if sub_palavra.isupper() and len(sub_palavra) > 1:
                    resultado.append(sub_palavra)
                else:
                    resultado.append(sub_palavra.capitalize())
            elif sub_palavra.lower() in preposicoes:
                resultado.append(sub_palavra.lower())
            else:
                if sub_palavra.isupper() and len(sub_palavra) > 1:
                    resultado.append(sub_palavra)
                else:
                    resultado.append(sub_palavra.capitalize())
    return " ".join(resultado)

def formatar_nome_final(nome, remove_spaces, camel_case):
    if not nome:
        return ""
    if camel_case:
        palavras = nome.split()
        if len(palavras) == 1:
            w = palavras[0]
            if w.isupper():
                nome = w.lower()
            else:
                nome = w[0].lower() + w[1:] if w else ""
        elif len(palavras) > 1:
            primeira = palavras[0].lower()
            resto = [p.capitalize() for p in palavras[1:]]
            nome = primeira + "".join(resto)
    elif remove_spaces:
        nome = nome.replace(" ", "")
    return nome

def extrair_titulo_do_texto(texto):
    if not texto:
        return None
    linhas = [l.strip() for l in texto.split('\n') if l.strip()]
    palavras_descarte = {
        'issn', 'doi:', 'http', 'vol.', 'n.', 'jan.', 'jun.', 'recebido em', 
        'reformulado em', 'revista', 'universidade', 'editorial', 'artigo',
        'submetido', 'aceito', 'publicado', 'anais', 'proceedings', 'page', 'página',
        'blind review', 'diretrizes', 'licença', 'creative commons', 'todos os direitos'
    }
    
    linhas_filtradas = []
    for linha in linhas:
        if re.match(r'^\d+$', linha) or len(linha) < 10:
            continue
        if any(p in linha.lower() for p in palavras_descarte):
            continue
        linhas_filtradas.append(linha)
        
    titulo_partes = []
    for line in linhas_filtradas:
        if any(p in line.lower() for p in ['autores', 'authors', 'effects of', 'efectos de', 'resumo', 'abstract', 'resumen']):
            break
        titulo_partes.append(line)
        if len(titulo_partes) >= 3:
            break
    return " ".join(titulo_partes) if titulo_partes else None

def obter_titulo_pdf(caminho):
    try:
        reader = pypdf.PdfReader(caminho)
        meta = reader.metadata
        titulo = None
        if meta and meta.title:
            t_meta = meta.title.strip()
            titulos_genericos = {
                '', 'untitled', 'document', 'documento', ':: plataforma a', 'plataforma a',
                'capítulo do livro', 'capítulo de livro', 'capitulo do livro', 'capitulo de livro',
                'microsoft word'
            }
            if t_meta.lower() not in titulos_genericos:
                titulo = t_meta
        if not titulo and reader.pages:
            texto_p1 = reader.pages[0].extract_text()
            titulo = extrair_titulo_do_texto(texto_p1)
        return titulo
    except Exception:
        pass
    return None

def obter_titulo_docx(caminho):
    try:
        doc = docx.Document(caminho)
        for p in doc.paragraphs:
            texto = p.text.strip()
            if texto:
                if len(texto) > 150:
                    texto = texto[:147] + "..."
                return texto
    except Exception:
        pass
    return None

def scan_diretorio_backend(caminho_base, caminho_atual=None, prefixo_pai="", mode="content", query="", remove_spaces=False, camel_case=False):
    if caminho_atual is None:
        caminho_atual = Path(caminho_base)
        
    if caminho_atual == Path(caminho_base):
        prefixo_atual = ""
    else:
        num = obter_numero_diretorio(caminho_atual.name)
        if num:
            prefixo_atual = f"{prefixo_pai}.{num}" if prefixo_pai else num
        else:
            prefixo_atual = prefixo_pai
            
    files = []
    
    # Listar itens ordenados
    try:
        itens = sorted(list(caminho_atual.iterdir()), key=lambda x: x.name)
    except Exception:
        return []

    arquivos = []
    subdirs = []
    for item in itens:
        if item.name.startswith('.'):
            continue
        if item.name in ['index.html', 'renomear_arquivos.py', 'renomear_por_conteudo.py', 'servidor.py']:
            continue
        if item.is_file():
            arquivos.append(item)
        elif item.is_dir():
            subdirs.append(item)

    nomes_na_pasta = set()
    for idx, item in enumerate(arquivos, start=1):
        ext = item.suffix.lower()
        
        # Se houver busca por termo, verifica se o termo está no conteúdo do arquivo
        if query:
            termo = query.lower()
            conteudo_contem_termo = False
            try:
                if ext == '.pdf':
                    reader = pypdf.PdfReader(item)
                    for page in reader.pages[:10]: # busca nas primeiras 10 páginas
                        text = page.extract_text()
                        if text and termo in text.lower():
                            conteudo_contem_termo = True
                            break
                elif ext == '.docx':
                    doc = docx.Document(item)
                    for p in doc.paragraphs:
                        if termo in p.text.lower():
                            conteudo_contem_termo = True
                            break
                else:
                    # arquivos de texto normais
                    if ext in ['.txt', '.md', '.json', '.html', '.css', '.js']:
                        with open(item, 'r', encoding='utf-8', errors='ignore') as f:
                            if termo in f.read().lower():
                                conteudo_contem_termo = True
            except Exception:
                pass
                
            if not conteudo_contem_termo:
                continue # pula este arquivo se não contiver o termo de busca
                
        nome_limpo = re.sub(r'^\d+(\.\d+)*[_\s.-]+', '', item.name)
        
        if mode == 'number_only':
            if prefixo_atual:
                novo_nome = f"{prefixo_atual}.{idx}{ext}"
            else:
                novo_nome = f"{idx}{ext}"
        else:
            titulo = None
            if mode == 'content':
                if ext == '.pdf':
                    titulo = obter_titulo_pdf(item)
                elif ext == '.docx':
                    titulo = obter_titulo_docx(item)
            
            if titulo:
                nome_padronizado = limpar_titulo_para_arquivo(titulo)
            else:
                stem = Path(nome_limpo).stem
                nome_padronizado = limpar_titulo_para_arquivo(stem)
                
            nome_padronizado = formatar_nome_final(nome_padronizado, remove_spaces, camel_case) + ext
                
            if mode == 'number_name':
                if prefixo_atual:
                    prefixo_completo = f"{prefixo_atual}.{idx}"
                else:
                    prefixo_completo = str(idx)
                
                if remove_spaces or camel_case:
                    novo_nome = f"{prefixo_completo}.{nome_padronizado}"
                else:
                    novo_nome = f"{prefixo_completo}. {nome_padronizado}"
            else:
                if prefixo_atual:
                    if remove_spaces or camel_case:
                        novo_nome = f"{prefixo_atual}.{nome_padronizado}"
                    else:
                        novo_nome = f"{prefixo_atual}. {nome_padronizado}"
                else:
                    novo_nome = nome_padronizado
                
        # Resolver colisão de nomes na mesma pasta
        nome_proposto = novo_nome
        contador = 1
        while True:
            # Verifica se o nome já foi proposto nesta sessão
            if nome_proposto in nomes_na_pasta:
                colisao = True
            else:
                # Verifica se o arquivo já existe fisicamente e não é o próprio arquivo atual
                path_proposto = caminho_atual / nome_proposto
                if path_proposto.exists() and path_proposto.resolve() != item.resolve():
                    colisao = True
                else:
                    colisao = False
                    
            if not colisao:
                break
                
            # Se colidir, adiciona um sufixo numérico (ex: "1.1. Nome (2).pdf")
            stem_p = Path(novo_nome).stem
            ext_p = Path(novo_nome).suffix
            # Remove qualquer sufixo de colisão anterior para não acumular
            stem_p = re.sub(r'\s\(\d+\)$', '', stem_p)
            nome_proposto = f"{stem_p} ({contador + 1}){ext_p}"
            contador += 1
            
        novo_nome = nome_proposto
        nomes_na_pasta.add(novo_nome)
        
        files.append({
            'originalName': item.name,
            'newName': novo_nome,
            'relativePath': str(item.relative_to(caminho_base)),
            'absolutePath': str(item.resolve()),
            'status': 'ready' if item.name != novo_nome else 'success'
        })

    for subdir in subdirs:
        files.extend(scan_diretorio_backend(caminho_base, subdir, prefixo_atual, mode, query, remove_spaces, camel_case))
            
    return files

class APIHandler(http.server.SimpleHTTPRequestHandler):
    def end_headers(self):
        self.send_header('Access-Control-Allow-Origin', '*')
        self.send_header('Access-Control-Allow-Methods', 'GET, POST, OPTIONS')
        self.send_header('Access-Control-Allow-Headers', 'Content-Type')
        super().end_headers()

    def do_OPTIONS(self):
        self.send_response(200, "ok")
        self.end_headers()

    def do_GET(self):
        url = urllib.parse.urlparse(self.path)
        if url.path == '/api/select_folder':
            try:
                import subprocess
                cmd = ["osascript", "-e", 'POSIX path of (choose folder with prompt "Selecione a Pasta para Organizar")']
                resultado = subprocess.run(cmd, capture_output=True, text=True)
                pasta_selecionada = resultado.stdout.strip()
                
                self.send_response(200)
                self.send_header('Content-Type', 'application/json')
                self.end_headers()
                self.wfile.write(json.dumps({'path': pasta_selecionada}).encode())
            except Exception as e:
                self.send_response(500)
                self.send_header('Content-Type', 'application/json')
                self.end_headers()
                self.wfile.write(json.dumps({'error': str(e)}).encode())
        elif url.path == '/api/scan':
            query = urllib.parse.parse_qs(url.query)
            caminho_base = query.get('path', [''])[0]
            mode = query.get('mode', ['content'])[0]
            q = query.get('q', [''])[0]
            remove_spaces = query.get('remove_spaces', ['false'])[0].lower() == 'true'
            camel_case = query.get('camel_case', ['false'])[0].lower() == 'true'
            
            if not caminho_base:
                caminho_base = "."
                
            caminho_abs = Path(caminho_base).resolve()
            if not caminho_abs.exists():
                self.send_response(400)
                self.send_header('Content-Type', 'application/json')
                self.end_headers()
                self.wfile.write(json.dumps({'error': 'Diretório não existe'}).encode())
                return
                
            files = scan_diretorio_backend(caminho_abs, mode=mode, query=q, remove_spaces=remove_spaces, camel_case=camel_case)
            self.send_response(200)
            self.send_header('Content-Type', 'application/json')
            self.end_headers()
            self.wfile.write(json.dumps({'files': files, 'resolvedPath': str(caminho_abs)}).encode())
        else:
            # Servir arquivos estáticos (index.html, etc.) resolvendo caminhos do PyInstaller se compilado
            base_dir = getattr(sys, '_MEIPASS', os.path.dirname(os.path.abspath(__file__)))
            file_path = os.path.join(base_dir, self.path.lstrip('/'))
            if self.path == '/':
                file_path = os.path.join(base_dir, 'index.html')
                
            if os.path.exists(file_path) and os.path.isfile(file_path):
                self.send_response(200)
                if file_path.endswith('.html'):
                    self.send_header('Content-Type', 'text/html; charset=utf-8')
                elif file_path.endswith('.js'):
                    self.send_header('Content-Type', 'application/javascript')
                elif file_path.endswith('.css'):
                    self.send_header('Content-Type', 'text/css')
                self.end_headers()
                with open(file_path, 'rb') as f:
                    self.wfile.write(f.read())
            else:
                self.send_response(404)
                self.end_headers()

    def do_POST(self):
        global historico_desfazer
        url = urllib.parse.urlparse(self.path)
        if url.path == '/api/rename':
            content_length = int(self.headers['Content-Length'])
            post_data = self.rfile.read(content_length)
            dados = json.loads(post_data.decode('utf-8'))
            
            caminho_base = dados.get('path', '')
            alteracoes = dados.get('changes', [])
            
            resultados = []
            for alt in alteracoes:
                abs_path = Path(alt['absolutePath'])
                novo_nome = alt['newName']
                novo_path = abs_path.parent / novo_nome
                
                try:
                    if abs_path.exists() and abs_path != novo_path:
                        if novo_path.exists():
                            resultados.append({'absolutePath': alt['absolutePath'], 'status': 'error', 'errorMsg': 'O arquivo de destino já existe. Operação cancelada para não sobrescrever.'})
                        else:
                            abs_path.rename(novo_path)
                            resultados.append({'absolutePath': alt['absolutePath'], 'status': 'success', 'newName': novo_nome})
                    else:
                        resultados.append({'absolutePath': alt['absolutePath'], 'status': 'success', 'newName': alt['originalName']})
                except Exception as e:
                    resultados.append({'absolutePath': alt['absolutePath'], 'status': 'error', 'errorMsg': str(e)})
                    
            # Armazenar no histórico de desfazer (se houver mudanças reais e bem-sucedidas)
            lote_desfazer = []
            for r in resultados:
                if r['status'] == 'success' and r.get('newName') and r['newName'] != Path(r['absolutePath']).name:
                    lote_desfazer.append((str(Path(r['absolutePath']).parent / r['newName']), r['absolutePath']))
            if lote_desfazer:
                historico_desfazer.append(lote_desfazer)
                
            self.send_response(200)
            self.send_header('Content-Type', 'application/json')
            self.end_headers()
            self.wfile.write(json.dumps({'results': resultados}).encode())
            
        elif url.path == '/api/undo':
            if not historico_desfazer:
                self.send_response(200)
                self.send_header('Content-Type', 'application/json')
                self.end_headers()
                self.wfile.write(json.dumps({'status': 'empty', 'message': 'Nenhum histórico para desfazer.'}).encode())
                return
                
            ultimo_lote = historico_desfazer.pop()
            resultados = []
            for novo_caminho, antigo_caminho in ultimo_lote:
                novo_p = Path(novo_caminho)
                antigo_p = Path(antigo_caminho)
                try:
                    if novo_p.exists() and not antigo_p.exists():
                        novo_p.rename(antigo_p)
                        resultados.append({'novoCaminho': novo_caminho, 'antigoCaminho': antigo_caminho, 'status': 'success'})
                    else:
                        resultados.append({'novoCaminho': novo_caminho, 'antigoCaminho': antigo_caminho, 'status': 'error', 'errorMsg': 'Arquivo já existe ou original sumiu.'})
                except Exception as e:
                    resultados.append({'novoCaminho': novo_caminho, 'antigoCaminho': antigo_caminho, 'status': 'error', 'errorMsg': str(e)})
                    
            self.send_response(200)
            self.send_header('Content-Type', 'application/json')
            self.end_headers()
            self.wfile.write(json.dumps({'status': 'done', 'results': resultados}).encode())

# Executar servidor e janela desktop
import threading
import sys
import webview

def rodar_servidor():
    socketserver.TCPServer.allow_reuse_address = True
    with socketserver.TCPServer(("", PORT), APIHandler) as httpd:
        print(f"Servidor API local rodando na porta {PORT}...")
        httpd.serve_forever()

if __name__ == "__main__":
    # Inicia o servidor HTTP em uma thread secundária
    t = threading.Thread(target=rodar_servidor, daemon=True)
    t.start()
    
    # Inicia a janela desktop nativa usando webview
    webview.create_window("Organizador de Arquivos Inteligente", f"http://localhost:{PORT}", width=1280, height=800)
    webview.start()
